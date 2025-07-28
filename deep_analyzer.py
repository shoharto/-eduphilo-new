import docx
import os
from pathlib import Path

def deep_analyze_docx(docx_path):
    """
    Deep analysis of .docx file to understand its structure
    """
    try:
        if not os.path.exists(docx_path):
            print(f"Error: File '{docx_path}' not found.")
            return None
        
        doc = docx.Document(docx_path)
        
        print("=== DEEP DOCUMENT ANALYSIS ===")
        print(f"Document has {len(doc.paragraphs)} paragraphs")
        print(f"Document has {len(doc.tables)} tables")
        print(f"Document has {len(doc.sections)} sections")
        print(f"Document has {len(doc.styles)} styles")
        
        # Analyze each paragraph in detail
        all_content = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            print(f"\n--- Paragraph {i+1} ---")
            print(f"Style: {paragraph.style.name}")
            print(f"Text: '{paragraph.text}'")
            print(f"Number of runs: {len(paragraph.runs)}")
            
            # Analyze each run in the paragraph
            for j, run in enumerate(paragraph.runs):
                print(f"  Run {j+1}: '{run.text}'")
                print(f"    Bold: {run.bold}")
                print(f"    Italic: {run.italic}")
                print(f"    Font size: {run.font.size}")
                print(f"    Font name: {run.font.name}")
                
                if run.text.strip():
                    all_content.append(run.text.strip())
        
        # Analyze tables
        for i, table in enumerate(doc.tables):
            print(f"\n--- Table {i+1} ---")
            print(f"Rows: {len(table.rows)}")
            print(f"Columns: {len(table.columns)}")
            
            for row_idx, row in enumerate(table.rows):
                print(f"  Row {row_idx+1}:")
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    print(f"    Cell {col_idx+1}: '{cell_text}'")
                    if cell_text:
                        all_content.append(cell_text)
        
        # Analyze styles
        print(f"\n--- Available Styles ---")
        for style in doc.styles:
            print(f"  {style.name}")
        
        return all_content
        
    except Exception as e:
        print(f"Error analyzing file: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def extract_raw_text(docx_path):
    """
    Extract raw text without any processing
    """
    try:
        doc = docx.Document(docx_path)
        raw_text = []
        
        for paragraph in doc.paragraphs:
            raw_text.append(paragraph.text)
        
        return '\n'.join(raw_text)
        
    except Exception as e:
        print(f"Error extracting raw text: {str(e)}")
        return None

def main():
    docx_file = "eduphilo-website-requirements.docx"
    
    print("Starting deep analysis...")
    content = deep_analyze_docx(docx_file)
    
    print(f"\n=== RAW TEXT EXTRACTION ===")
    raw_text = extract_raw_text(docx_file)
    
    if raw_text:
        print(f"Raw text length: {len(raw_text)} characters")
        print("Raw text content:")
        print("-" * 50)
        print(raw_text)
        print("-" * 50)
        
        # Save raw text
        with open("raw_content.txt", "w", encoding="utf-8") as f:
            f.write(raw_text)
        print("Raw content saved to 'raw_content.txt'")
    else:
        print("No raw text found!")
    
    if content:
        print(f"\nTotal content items: {len(content)}")
        print("Content items:")
        for i, item in enumerate(content):
            print(f"  {i+1}. {item}")

if __name__ == "__main__":
    main() 