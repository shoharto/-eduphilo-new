import docx
import os
import re
from pathlib import Path

def extract_all_content_from_docx(docx_path):
    """
    Extract all possible content from a .docx file
    """
    try:
        if not os.path.exists(docx_path):
            print(f"Error: File '{docx_path}' not found.")
            return None
        
        doc = docx.Document(docx_path)
        
        print(f"Document has {len(doc.paragraphs)} paragraphs")
        print(f"Document has {len(doc.tables)} tables")
        print(f"Document has {len(doc.sections)} sections")
        
        # Extract all text content
        all_content = []
        
        # Process all paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"Paragraph {i+1}: {text[:100]}...")
                all_content.append(text)
        
        # Process all tables
        for i, table in enumerate(doc.tables):
            print(f"Table {i+1}:")
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    print(f"  Row {row_idx+1}: {' | '.join(row_text)}")
                    all_content.extend(row_text)
        
        return all_content
        
    except Exception as e:
        print(f"Error reading file: {str(e)}")
        return None

def convert_docx_to_markdown_improved(docx_path, output_path=None):
    """
    Improved conversion with better content extraction
    """
    try:
        if not os.path.exists(docx_path):
            print(f"Error: File '{docx_path}' not found.")
            return False
        
        doc = docx.Document(docx_path)
        
        md_content = []
        
        # Add document title
        md_content.append("# Eduphilo Website Requirements")
        md_content.append("")
        
        # Process paragraphs with better detection
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check for various heading patterns
            style_name = paragraph.style.name.lower()
            font_size = None
            
            # Try to get font size
            for run in paragraph.runs:
                if run.font.size:
                    font_size = run.font.size.pt
                    break
            
            # Determine if it's a heading based on style and font size
            if ('heading' in style_name or 
                'title' in style_name or 
                (font_size and font_size > 14) or
                paragraph.text.isupper() or
                len(text) < 100 and text.endswith(':')):
                
                # Determine heading level
                if 'heading 1' in style_name or font_size and font_size > 18:
                    md_content.append(f"# {text}")
                elif 'heading 2' in style_name or font_size and font_size > 16:
                    md_content.append(f"## {text}")
                elif 'heading 3' in style_name or font_size and font_size > 14:
                    md_content.append(f"### {text}")
                else:
                    md_content.append(f"## {text}")
            else:
                # Regular paragraph
                md_content.append(text)
            
            md_content.append("")
        
        # Process tables
        for table in doc.tables:
            if len(table.rows) > 0:
                md_content.append("")
                
                # Create table
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text if cell_text else " ")
                    
                    if row_data:
                        md_content.append("| " + " | ".join(row_data) + " |")
                        
                        # Add separator after header row
                        if row_idx == 0:
                            separator = "| " + " | ".join(["---"] * len(row_data)) + " |"
                            md_content.append(separator)
                
                md_content.append("")
        
        # Join content
        final_content = '\n'.join(md_content)
        
        # Determine output path
        if output_path is None:
            base_name = Path(docx_path).stem
            output_path = f"{base_name}.md"
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_content)
        
        print(f"Successfully converted '{docx_path}' to '{output_path}'")
        print(f"Content length: {len(final_content)} characters")
        
        return True
        
    except Exception as e:
        print(f"Error converting file: {str(e)}")
        return False

def main():
    docx_file = "eduphilo-website-requirements.docx"
    output_file = "eduphilo-website-requirements.md"
    
    print("=== DOCUMENT ANALYSIS ===")
    content = extract_all_content_from_docx(docx_file)
    
    if content:
        print(f"\nTotal content items found: {len(content)}")
        print("\n=== CONVERTING TO MARKDOWN ===")
        success = convert_docx_to_markdown_improved(docx_file, output_file)
        
        if success:
            print(f"\nConversion completed!")
            print(f"Output file: {output_file}")
            
            # Show preview
            try:
                with open(output_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                    if content:
                        print(f"\nPreview (first 1000 characters):")
                        print("-" * 50)
                        print(content[:1000] + "..." if len(content) > 1000 else content)
                        print("-" * 50)
                    else:
                        print("Warning: Generated file is empty!")
            except Exception as e:
                print(f"Error reading output file: {str(e)}")
        else:
            print("Conversion failed!")
    else:
        print("No content found in the document!")

if __name__ == "__main__":
    main() 