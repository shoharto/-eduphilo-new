import os
import sys
from pathlib import Path

def try_mammoth(docx_path):
    """
    Try to extract content using mammoth library
    """
    try:
        import mammoth
        
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
            
            # Convert HTML to plain text
            import re
            text_content = re.sub(r'<[^>]+>', '', html_content)
            text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
            
            return text_content.strip()
            
    except Exception as e:
        print(f"Mammoth failed: {str(e)}")
        return None

def try_textract(docx_path):
    """
    Try to extract content using textract library
    """
    try:
        import textract
        text = textract.process(docx_path).decode('utf-8')
        return text.strip()
        
    except Exception as e:
        print(f"Textract failed: {str(e)}")
        return None

def try_docx2txt(docx_path):
    """
    Try to extract content using docx2txt library
    """
    try:
        import docx2txt
        text = docx2txt.process(docx_path)
        return text.strip()
        
    except Exception as e:
        print(f"Docx2txt failed: {str(e)}")
        return None

def try_python_docx_enhanced(docx_path):
    """
    Enhanced python-docx extraction
    """
    try:
        import docx
        
        doc = docx.Document(docx_path)
        content = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        content.append(cell.text.strip())
        
        # Extract from headers and footers
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
            
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    if paragraph.text.strip():
                        content.append(paragraph.text.strip())
        
        return '\n\n'.join(content)
        
    except Exception as e:
        print(f"Enhanced python-docx failed: {str(e)}")
        return None

def convert_to_markdown(content, output_path):
    """
    Convert extracted content to markdown format
    """
    if not content:
        return False
    
    # Basic markdown conversion
    lines = content.split('\n')
    md_lines = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Detect headings (simple heuristic)
        if line.isupper() and len(line) < 100:
            md_lines.append(f"# {line}")
        elif line.endswith(':') and len(line) < 100:
            md_lines.append(f"## {line}")
        elif line.startswith(('•', '-', '*', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            # Convert to markdown list
            line = line.replace('•', '-').replace('*', '-')
            md_lines.append(line)
        else:
            md_lines.append(line)
        
        md_lines.append("")  # Add spacing
    
    final_content = '\n'.join(md_lines)
    
    # Write to file
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(final_content)
    
    return True

def main():
    docx_file = "eduphilo-website-requirements.docx"
    output_file = "eduphilo-website-requirements.md"
    
    print("=== ROBUST DOCUMENT EXTRACTION ===")
    
    # Try different methods
    methods = [
        ("Mammoth", try_mammoth),
        ("Textract", try_textract),
        ("Docx2txt", try_docx2txt),
        ("Enhanced python-docx", try_python_docx_enhanced)
    ]
    
    extracted_content = None
    
    for method_name, method_func in methods:
        print(f"\nTrying {method_name}...")
        content = method_func(docx_file)
        
        if content and len(content.strip()) > 0:
            print(f"✓ {method_name} succeeded!")
            print(f"Content length: {len(content)} characters")
            print("Preview:")
            print("-" * 50)
            print(content[:500] + "..." if len(content) > 500 else content)
            print("-" * 50)
            
            extracted_content = content
            break
        else:
            print(f"✗ {method_name} failed or returned empty content")
    
    if extracted_content:
        print(f"\n=== CONVERTING TO MARKDOWN ===")
        success = convert_to_markdown(extracted_content, output_file)
        
        if success:
            print(f"✓ Successfully converted to '{output_file}'")
            
            # Show final result
            with open(output_file, 'r', encoding='utf-8') as f:
                final_content = f.read()
                print(f"Final markdown length: {len(final_content)} characters")
                print("\nFinal markdown preview:")
                print("-" * 50)
                print(final_content[:1000] + "..." if len(final_content) > 1000 else final_content)
                print("-" * 50)
        else:
            print("✗ Markdown conversion failed")
    else:
        print("\n✗ All extraction methods failed!")
        print("The document might be:")
        print("- Empty or corrupted")
        print("- Password protected")
        print("- Contains only images or embedded objects")
        print("- In a different format")

if __name__ == "__main__":
    main() 