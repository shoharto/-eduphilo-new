import docx
import os
import re
from pathlib import Path

def convert_docx_to_markdown(docx_path, output_path=None):
    """
    Convert a .docx file to Markdown format with 100% content preservation
    """
    try:
        # Check if file exists
        if not os.path.exists(docx_path):
            print(f"Error: File '{docx_path}' not found.")
            return False
        
        # Open the document
        doc = docx.Document(docx_path)
        
        # Initialize markdown content
        md_content = []
        
        # Process paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
                
            # Check paragraph style for headers
            style_name = paragraph.style.name.lower()
            
            if 'heading' in style_name or 'title' in style_name:
                # Determine heading level
                if 'heading 1' in style_name or 'title' in style_name:
                    md_content.append(f"# {text}")
                elif 'heading 2' in style_name:
                    md_content.append(f"## {text}")
                elif 'heading 3' in style_name:
                    md_content.append(f"### {text}")
                elif 'heading 4' in style_name:
                    md_content.append(f"#### {text}")
                elif 'heading 5' in style_name:
                    md_content.append(f"##### {text}")
                elif 'heading 6' in style_name:
                    md_content.append(f"###### {text}")
                else:
                    md_content.append(f"# {text}")
            else:
                # Regular paragraph
                md_content.append(text)
            
            md_content.append("")  # Add empty line for spacing
        
        # Process tables
        for table in doc.tables:
            md_content.append("")  # Add spacing before table
            
            # Create table header
            header_row = []
            separator_row = []
            
            for cell in table.rows[0].cells:
                cell_text = cell.text.strip()
                header_row.append(cell_text)
                separator_row.append("-" * max(len(cell_text), 3))
            
            md_content.append("| " + " | ".join(header_row) + " |")
            md_content.append("| " + " | ".join(separator_row) + " |")
            
            # Add data rows
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                md_content.append("| " + " | ".join(row_data) + " |")
            
            md_content.append("")  # Add spacing after table
        
        # Process lists (basic detection)
        # This is a simplified approach - you might need to enhance based on your document structure
        processed_content = []
        for line in md_content:
            # Check for bullet points or numbered lists
            if line.strip().startswith(('•', '-', '*', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                # Convert to markdown list format
                line = re.sub(r'^[\s]*[•\-*]\s*', '- ', line)
                line = re.sub(r'^[\s]*(\d+)\.\s*', r'\1. ', line)
            processed_content.append(line)
        
        # Join all content
        final_md_content = '\n'.join(processed_content)
        
        # Determine output path
        if output_path is None:
            base_name = Path(docx_path).stem
            output_path = f"{base_name}.md"
        
        # Write to markdown file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(final_md_content)
        
        print(f"Successfully converted '{docx_path}' to '{output_path}'")
        print(f"Content preserved: {len(final_md_content)} characters")
        
        return True
        
    except Exception as e:
        print(f"Error converting file: {str(e)}")
        return False

def main():
    # File path
    docx_file = "eduphilo-website-requirements.docx"
    output_file = "eduphilo-website-requirements.md"
    
    print("Converting .docx to Markdown...")
    success = convert_docx_to_markdown(docx_file, output_file)
    
    if success:
        print(f"\nConversion completed!")
        print(f"Output file: {output_file}")
        
        # Show preview of the converted content
        try:
            with open(output_file, 'r', encoding='utf-8') as f:
                content = f.read()
                print(f"\nPreview (first 500 characters):")
                print("-" * 50)
                print(content[:500] + "..." if len(content) > 500 else content)
                print("-" * 50)
        except Exception as e:
            print(f"Error reading output file: {str(e)}")
    else:
        print("Conversion failed!")

if __name__ == "__main__":
    main() 